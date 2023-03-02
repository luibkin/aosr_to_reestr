import docx
import pandas as pd
import os


# Функция определения строк (rows), содержащих освидетельствуемые работы
def job_rows(a, f):
    k = 0
    m = 0
    doc = docx.Document('aosr/' + f)
    for i in a:
        x = doc.tables[2].rows[i].cells[0].text
        if x == '\nи составили настоящий акт о нижеследующем:' or x == 'и составили настоящий акт о нижеследующем:':
            k = i + 1
            print(k)
        if x == '(наименование скрытых работ)':
            m = i
            print(m)
    return [k, m]


all_data = {}
k = 0
r = 3
date_data = {}
final_data = {}

files = os.listdir('aosr')

for file in files:
    print(file)
    doc = docx.Document('aosr/' + file)
    date = '20' + doc.tables[1].rows[0].cells[7].text + '.' + doc.tables[1].rows[0].cells[5].text + '.' + doc.tables[1].rows[0].cells[3].text
    date_1 = doc.tables[1].rows[0].cells[3].text + '.' + doc.tables[1].rows[0].cells[5].text+ '.' + '20' + doc.tables[1].rows[0].cells[7].text
#    print(doc.tables[1].rows[0].cells[3].text)
#    print(doc.tables[1].rows[0].cells[5].text)
#    print(doc.tables[1].rows[0].cells[7].text)
#    print(doc.tables[2].rows[4].cells[0].text) # работы
    j = job_rows(range(20, 40), file)
#     print(j)
    jobs_num = 0
    jobs = ''
    for p in range(j[0], j[1]):
#        print(doc.tables[2].rows[p].cells[0].text)
#        print(len(range(j[0], j[1])))
        if len(range(j[0], j[1])) == 1:
            jobs = doc.tables[2].rows[p].cells[0].text
            if jobs == ' ' or jobs == '': 
                jobs =  doc.tables[2].rows[p].cells[1].text
        else:
            jobs = jobs + doc.tables[2].rows[p].cells[0].text
            jobs_num = jobs_num + 1
#    print(jobs, 'jobs_num = ', jobs_num, '\n')
# Формирование словаря из ячеек с необходимыми данными.
# [5, 0], [7, 0], [9, 14], [27, 10] - координаты ячеек соответственно:
# номер акта, материал, дата, сертификаты. Можно добавить еще ячейки.
# Ячейка с сертификатом в данном случае в дальнейшем не используется.
    all_data[k] = {0: doc.tables[1].rows[0].cells[1].text, 1: date, 4: jobs, 5: date_1}

# Формирование словаря из ячеек с датой (для упорядочивания результата
# по дате).
    date_data[k] = date
    k = k + 1
#print(all_data)
# Получение списка ключей словаря из ячеек с датой при упорядочивании словаря
# по значениям даты.
#for date in date_data:
#    date_data[p] = date_data[p][6:10] + '.' + date_data[p][3:5] + '.' + date_data[p][:2]
#    p = p + 1
sorted_date_data_keys = sorted(date_data, key=date_data.get)  # [1, 3, 2] - пример результата.
#print(sorted_date_data_keys)
#print([date_data[k] for k in sorted_date_data_keys])
# Формирование словаря из ячеек с необходимыми данными под запись в реестр
# (с упорядочиванием по дате).
for key_b in sorted_date_data_keys:
#    final_data[key_b] = {0: [('Акт освидетельствования скрытых работ №' + str(all_data[key_b][0]) + ' от ' + str(all_data[key_b][5]) + 'г.')], 1: (' (' + all_data[key_b][4] + ').')}
    final_data[key_b] = {0: [('Акт освидетельствования скрытых работ' + ' от ' + str(all_data[key_b][5]) + 'г.' + ' (' + str(all_data[key_b][4]) + ').')], 1: ('№' + str(all_data[key_b][0])), 2: str(all_data[key_b][5])}
    for value in final_data[key_b].values():
        print(value)

# Запись в реестр.
with pd.ExcelWriter('./1.xlsx', engine='xlsxwriter') as writer:
    for key_c in sorted_date_data_keys:
        df2 = pd.DataFrame(final_data[key_c])
        df2.to_excel(writer, sheet_name="Sheet1", startrow=r, header=False, index=False)
        r = r + 1
#    print(final_data[key_b])
"""r = 0

filename = 'aosr/2.docx'

doc = docx.Document(filename)

print(len(doc.tables))

print(type(doc.tables[1].rows[0].cells[0].text))

print(doc.tables[1].rows[0].cells[1].text) # номер акта

print(doc.tables[1].rows[0].cells[3].text)

print(doc.tables[1].rows[0].cells[5].text)

print(doc.tables[1].rows[0].cells[7].text)

print(doc.tables[2].rows[28].cells[1].text) # работы"""

"""for i in doc.paragraphs:
    print(doc.paragraphs[r].text)
    r = r + 1"""




"""import os

# Поиск по папкам
paths = []
folder = os.getcwd()
for root, dirs, files in os.walk(folder):
    for file in files:
        if file.endswith('docx') and not file.startswith('~'):
            paths.append(os.path.join(root, file))"""